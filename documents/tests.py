"""
Test suite. Runs entirely offline - no Groq key, no network.

    python manage.py test -v 1

Every check here asserts something that would be a real defect: a signature
going undetected, a vector chart being invisible to search, an ATS-unsafe
export, or a score that cannot be explained.
"""

import os
import tempfile

from django.contrib.auth.models import User
from django.core.files.base import ContentFile
from django.test import TestCase, override_settings

try:
    import pymupdf
except ImportError:  # pragma: no cover
    import fitz as pymupdf

from documents import retrieval
from documents.chunking import chunk_pages, clean
from documents.models import Chunk, Document, Figure
from documents.services import process_document
from resume import ats, export, parsing

MEDIA = tempfile.mkdtemp(prefix="docqa-test-")


# --------------------------------------------------------------- fixtures

def build_test_pdf() -> bytes:
    """A PDF containing every kind of thing the extractor must find."""
    doc = pymupdf.open()

    # page 1 - plain prose
    p1 = doc.new_page()
    p1.insert_text((60, 90), "Quarterly Operations Report", fontsize=18)
    p1.insert_textbox(
        pymupdf.Rect(60, 120, 520, 400),
        "The distribution centre processed 48,120 shipments during the quarter, "
        "an increase of eleven percent. Cold chain compliance held at 99.2 percent "
        "across all lanes. The Karachi hub remains the constraint on throughput, "
        "and the mezzanine expansion is scheduled for the third quarter.\n\n"
        "Refund policy: claims must be submitted within thirty days of delivery, "
        "supported by photographic evidence and the original consignment note.",
        fontsize=11,
    )

    # page 2 - a vector bar chart, which carries no image object at all
    p2 = doc.new_page()
    p2.insert_text((60, 70), "Figure 1: Monthly shipment volume", fontsize=11)
    base_y = 360
    for i, h in enumerate([70, 95, 60, 130, 110, 145]):
        x = 80 + i * 55
        p2.draw_rect(pymupdf.Rect(x, base_y - h, x + 34, base_y),
                     color=(0.15, 0.3, 0.6), fill=(0.25, 0.45, 0.75))
        p2.insert_text((x + 6, base_y + 14), f"M{i + 1}", fontsize=8)
    p2.draw_line(pymupdf.Point(70, base_y), pymupdf.Point(420, base_y))
    p2.draw_line(pymupdf.Point(70, base_y), pymupdf.Point(70, 190))
    p2.insert_text((60, 430), "Volumes rose steadily through the period.", fontsize=10)

    # page 3 - a signature block
    p3 = doc.new_page()
    p3.insert_text((60, 90), "Authorisation", fontsize=14)
    p3.insert_textbox(
        pymupdf.Rect(60, 110, 520, 200),
        "The undersigned confirms that the figures presented above are accurate "
        "and complete to the best of their knowledge.",
        fontsize=11,
    )
    p3.insert_text((60, 300), "Signature:", fontsize=11)
    # a loose handwritten-looking stroke: wide, sparse, few colours
    pts = [(140, 300), (165, 275), (190, 315), (215, 272), (245, 308),
           (270, 280), (300, 305), (330, 285), (355, 300)]
    for a, b in zip(pts, pts[1:]):
        p3.draw_line(pymupdf.Point(*a), pymupdf.Point(*b), width=1.4)
    p3.insert_text((60, 340), "Date: 14 March 2026", fontsize=11)

    return doc.tobytes()


@override_settings(MEDIA_ROOT=MEDIA, LLM_BACKEND="echo", GROQ_API_KEY="",
                   DETECT_FIGURES=True, USE_RERANK=False)
class PdfEngineTests(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.user = User.objects.create_user("tester", password="x")

    def _ingest(self) -> Document:
        doc = Document(owner=self.user, title="report.pdf")
        doc.file.save("report.pdf", ContentFile(build_test_pdf()), save=True)
        process_document(doc.id)
        doc.refresh_from_db()
        return doc

    def test_document_becomes_ready(self):
        doc = self._ingest()
        self.assertEqual(doc.status, Document.Status.READY, doc.error)
        self.assertEqual(doc.page_count, 3)
        self.assertEqual(doc.pages.count(), 3)
        self.assertGreater(doc.num_chunks, 0)

    def test_pages_keep_their_numbers(self):
        doc = self._ingest()
        self.assertIn("Figure 1", doc.pages.get(number=2).text)

    def test_vector_chart_is_detected(self):
        """A bar chart drawn with rectangles has no image object.

        An extractor that only walks embedded images returns nothing here,
        which is why charts routinely go missing from document search.
        """
        doc = self._ingest()
        self.assertTrue(
            doc.figures.filter(source="vector", page_number=2).exists(),
            "vector chart on page 2 was not found",
        )

    def test_signature_block_is_detected(self):
        doc = self._ingest()
        self.assertTrue(
            doc.figures.filter(page_number=3).exists(),
            "nothing found in the signature block",
        )

    def test_figures_are_searchable(self):
        """Figures must become chunks, or they can never be retrieved."""
        doc = self._ingest()
        if doc.figures.filter(is_decorative=False).exists():
            figure_chunks = doc.chunks.filter(kind=Chunk.Kind.FIGURE)
            self.assertTrue(figure_chunks.exists())
            self.assertIsNotNone(figure_chunks.first().page_number)

    def test_figure_crops_are_stored(self):
        doc = self._ingest()
        for fig in doc.figures.all():
            self.assertTrue(fig.image, f"no crop saved for page {fig.page_number}")
            self.assertGreater(fig.image.size, 100)

    def test_reprocessing_does_not_duplicate(self):
        doc = self._ingest()
        first = doc.num_chunks
        process_document(doc.id)
        doc.refresh_from_db()
        self.assertEqual(doc.num_chunks, first)
        self.assertEqual(doc.chunks.count(), first)

    def test_encrypted_pdf_fails_with_a_readable_message(self):
        src = pymupdf.open("pdf", build_test_pdf())
        buf = src.tobytes(encryption=pymupdf.PDF_ENCRYPT_AES_256,
                          owner_pw="owner", user_pw="user")
        doc = Document(owner=self.user, title="locked.pdf")
        doc.file.save("locked.pdf", ContentFile(buf), save=True)
        process_document(doc.id)
        doc.refresh_from_db()
        self.assertEqual(doc.status, Document.Status.FAILED)
        self.assertIn("password", doc.error.lower())


@override_settings(MEDIA_ROOT=MEDIA, LLM_BACKEND="echo", GROQ_API_KEY="",
                   USE_RERANK=False)
class RetrievalTests(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.user = User.objects.create_user("searcher", password="x")
        cls.doc = Document.objects.create(
            owner=cls.user, title="handbook.txt", status=Document.Status.READY
        )
        passages = [
            "Refund claims must be submitted within thirty days of delivery.",
            "The Karachi hub is the constraint on overall throughput this quarter.",
            "Cold chain compliance held at 99.2 percent across all lanes.",
            "Annual leave accrues at two point five days per calendar month.",
            "Invoice AX-99213 was raised against the Lahore distribution account.",
        ]
        for i, text in enumerate(passages):
            Chunk.objects.create(
                document=cls.doc, index=i, text=text, page_number=i + 1,
                embedding=[0.0] * 384,
            )

    def _qs(self):
        return Chunk.objects.filter(document=self.doc)

    def test_bm25_finds_an_exact_identifier(self):
        """Vector search smears rare tokens; BM25 is what finds them."""
        hits = retrieval.search("AX-99213", self._qs(), top_k=3, use_rerank=False)
        self.assertTrue(hits)
        self.assertIn("AX-99213", hits[0].text)

    def test_bm25_ranks_the_right_passage_first(self):
        hits = retrieval.search("how long do I have to claim a refund",
                                self._qs(), top_k=3, use_rerank=False)
        self.assertTrue(hits)
        self.assertIn("thirty days", hits[0].text)

    def test_hits_carry_page_numbers(self):
        hits = retrieval.search("Karachi hub", self._qs(), top_k=2, use_rerank=False)
        self.assertTrue(hits)
        self.assertIsNotNone(hits[0].page_number)

    def test_no_match_returns_nothing_rather_than_noise(self):
        hits = retrieval.search("zzzqqq aardvark bewilderment",
                                self._qs(), top_k=5, use_rerank=False)
        self.assertEqual(hits, [])

    def test_rrf_rewards_agreement_between_rankers(self):
        """An item both rankers rank first must beat one only one ranker likes."""
        chunks = list(self._qs())[:4]
        agreed, lexical_only = chunks[0], chunks[1]
        fused = retrieval.reciprocal_rank_fusion(
            [agreed, lexical_only, chunks[2]],      # lexical ranking
            [agreed, chunks[3], chunks[2]],         # vector ranking
        )
        self.assertEqual(fused[0][0].id, agreed.id)
        scores = {c.id: s for c, s in fused}
        self.assertGreater(scores[agreed.id], scores[lexical_only.id])
        self.assertEqual(len(fused), 4)


class ChunkingTests(TestCase):
    def test_dehyphenation(self):
        self.assertIn("distribution", clean("distri-\nbution centre"))

    def test_chunks_carry_page_provenance(self):
        class P:
            def __init__(self, n, t):
                self.number, self.text = n, t

        pages = [P(1, "alpha " * 400), P(2, "beta " * 400)]
        chunks = list(chunk_pages(pages, target_tokens=80, overlap_tokens=10))
        self.assertTrue(chunks)
        self.assertTrue(all(c.page_start for c in chunks))
        self.assertTrue(any(c.page_start == 2 for c in chunks))

    def test_empty_input_yields_nothing(self):
        self.assertEqual(list(chunk_pages([])), [])


# ------------------------------------------------------------------ resume

CV_TEXT = """Ayesha Rahman
ayesha.rahman@example.com | +92 300 1234567 | Lahore, Pakistan
linkedin.com/in/ayesharahman

SUMMARY
Platform engineer with six years building and running container infrastructure.

EXPERIENCE
Senior Platform Engineer - Northwind Systems (2021 - Present)
- Migrated 40 services to Kubernetes, cutting deploy time from 45 minutes to 6
- Built CI pipelines in GitLab that run 900 jobs a day
- Responsible for on-call rota and incident response
- Reduced cloud spend by 32% through rightsizing and spot adoption

Platform Engineer - Bluewater Tech (2018 - 2021)
- Worked on monitoring and alerting
- Helped with database migrations

EDUCATION
BS Computer Science, University of Management and Technology, 2018

SKILLS
Python, Go, Kubernetes, Docker, Terraform, GitLab CI, AWS, PostgreSQL, Prometheus
"""

JOB_DESCRIPTION = """Senior Site Reliability Engineer

We are looking for an experienced SRE to own our Kubernetes platform.

Requirements:
- 5+ years of hands-on experience with Kubernetes in production
- Strong Terraform and infrastructure as code practice
- Must have experience with observability tooling, Prometheus and Grafana
- Proven track record reducing cloud cost
- Experience with GitOps workflows using ArgoCD is required
- Python or Go for automation

Nice to have:
- Service mesh experience, Istio preferred
- Experience running multi-region failover
"""


class ResumeParsingTests(TestCase):
    def setUp(self):
        os.makedirs(MEDIA, exist_ok=True)
        self.path = os.path.join(MEDIA, "cv.txt")
        with open(self.path, "w", encoding="utf-8") as fh:
            fh.write(CV_TEXT)
        self.parsed = parsing.parse(self.path, "cv.txt")

    def test_contact_is_extracted(self):
        c = self.parsed.contact
        self.assertEqual(c["email"], "ayesha.rahman@example.com")
        self.assertTrue(c["phone"])
        self.assertIn("ayesharahman", c["linkedin"])
        self.assertEqual(c["name"], "Ayesha Rahman")

    def test_sections_are_recognised(self):
        for section in ("summary", "experience", "education", "skills"):
            self.assertIn(section, self.parsed.sections)

    def test_skills_are_split(self):
        skills = [s.lower() for s in self.parsed.skills]
        self.assertIn("kubernetes", skills)
        self.assertIn("terraform", skills)

    def test_bullets_are_collected(self):
        self.assertGreaterEqual(len(self.parsed.bullets), 5)

    def test_weak_and_strong_openers_are_counted(self):
        self.assertGreater(self.parsed.stats["strong_openers"], 0)
        self.assertGreater(self.parsed.stats["weak_openers"], 0)  # Responsible, Worked

    def test_quantified_bullets_are_counted(self):
        self.assertGreaterEqual(self.parsed.stats["quantified_bullets"], 3)


class AtsFormatAuditTests(TestCase):
    def _pdf(self, two_column: bool, with_image: bool) -> str:
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_textbox(pymupdf.Rect(50, 60, 260, 700), CV_TEXT, fontsize=9)
        if two_column:
            page.insert_textbox(
                pymupdf.Rect(320, 60, 545, 700),
                "SKILLS\nKubernetes\nTerraform\nAWS\nPython\nGo\n"
                "Prometheus\nDocker\nGitLab\nPostgreSQL\nAnsible\nLinux",
                fontsize=9,
            )
        if with_image:
            pix = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 120, 120))
            pix.set_rect(pix.irect, (200, 120, 60))
            page.insert_image(pymupdf.Rect(430, 40, 545, 155), pixmap=pix)
        path = os.path.join(MEDIA, f"cv_{two_column}_{with_image}.pdf")
        doc.save(path)
        doc.close()
        return path

    def test_two_column_layout_is_flagged_as_critical(self):
        parsed = parsing.parse(self._pdf(True, False), "cv.pdf")
        codes = {i.code: i for i in parsed.issues}
        self.assertIn("multi_column", codes)
        self.assertEqual(codes["multi_column"].severity, "critical")
        self.assertTrue(codes["multi_column"].fix)

    def test_single_column_is_not_flagged(self):
        parsed = parsing.parse(self._pdf(False, False), "cv.pdf")
        self.assertNotIn("multi_column", {i.code for i in parsed.issues})

    def test_embedded_image_is_flagged(self):
        parsed = parsing.parse(self._pdf(False, True), "cv.pdf")
        self.assertIn("images", {i.code for i in parsed.issues})

    def test_docx_tables_are_flagged(self):
        from docx import Document as Docx

        d = Docx()
        d.add_paragraph("Ayesha Rahman")
        table = d.add_table(rows=3, cols=2)
        for row in table.rows:
            row.cells[0].text = "Kubernetes"
            row.cells[1].text = "6 years"
        path = os.path.join(MEDIA, "cv_table.docx")
        d.save(path)

        parsed = parsing.parse(path, "cv_table.docx")
        self.assertIn("tables", {i.code for i in parsed.issues})


class AtsScoringTests(TestCase):
    def setUp(self):
        os.makedirs(MEDIA, exist_ok=True)
        path = os.path.join(MEDIA, "cv_score.txt")
        with open(path, "w", encoding="utf-8") as fh:
            fh.write(CV_TEXT)
        self.parsed = parsing.parse(path, "cv_score.txt")
        self.reqs = ats.extract_requirements(JOB_DESCRIPTION)

    def test_requirements_are_extracted(self):
        self.assertTrue(self.reqs["must_have"])
        self.assertIn("kubernetes", self.reqs["keywords"])
        self.assertEqual(self.reqs["years_required"], 5)

    def test_score_is_explainable(self):
        result = ats.score(self.parsed, self.reqs)
        self.assertTrue(0 <= result["overall"] <= 100)
        self.assertEqual(
            set(result["components"]),
            {"format", "keywords", "content", "structure", "contact"},
        )
        for name, comp in result["components"].items():
            self.assertTrue(comp["notes"], f"{name} must explain itself")

    def test_score_is_deterministic(self):
        a = ats.score(self.parsed, self.reqs)
        b = ats.score(self.parsed, self.reqs)
        self.assertEqual(a["overall"], b["overall"])

    def test_missing_keywords_are_found_and_ranked(self):
        gaps = ats.gap_analysis(self.parsed, self.reqs)
        missing = {m["term"] for m in gaps["missing_keywords"]}
        self.assertTrue({"argocd", "grafana"} & missing, missing)
        self.assertNotIn("kubernetes", missing)
        self.assertNotIn("terraform", missing)

    def test_gap_analysis_gives_priorities(self):
        self.assertTrue(ats.gap_analysis(self.parsed, self.reqs)["priorities"])

    def test_no_job_description_still_scores(self):
        result = ats.score(self.parsed)
        self.assertFalse(result["assessed_against_job"])
        self.assertGreater(result["overall"], 0)


class ExportTests(TestCase):
    def setUp(self):
        os.makedirs(MEDIA, exist_ok=True)
        self.tailored = {
            "summary": "Platform engineer with six years running Kubernetes in production.",
            "skills": ["Kubernetes", "Terraform", "Prometheus", "Python", "Go"],
            "experience": [{
                "company": "Northwind Systems",
                "title": "Senior Platform Engineer",
                "dates": "2021 - Present",
                "bullets": [
                    "Migrated 40 services to Kubernetes, cutting deploy time from 45 minutes to 6",
                    "Reduced cloud spend by 32% through rightsizing and spot adoption",
                ],
            }],
            "projects": [],
            "education": ["BS Computer Science, UMT, 2018"],
            "certifications": [],
        }
        self.contact = {
            "name": "Ayesha Rahman",
            "email": "ayesha.rahman@example.com",
            "phone": "+92 300 1234567",
            "linkedin": "linkedin.com/in/ayesharahman",
        }

    def test_docx_is_produced(self):
        self.assertGreater(len(export.build_docx(self.tailored, self.contact)), 5000)

    def test_export_is_ats_safe_by_construction(self):
        """Prove it rather than assert it.

        If a future edit reintroduces a table or a header, this fails here
        instead of in somebody's job application.
        """
        report = export.verify_ats_safe(export.build_docx(self.tailored, self.contact))
        self.assertTrue(report["safe"], report["problems"])
        self.assertEqual(report["tables"], 0)

    def test_exported_docx_reparses_cleanly(self):
        path = os.path.join(MEDIA, "out.docx")
        with open(path, "wb") as fh:
            fh.write(export.build_docx(self.tailored, self.contact))

        parsed = parsing.parse(path, "out.docx")
        self.assertEqual(parsed.contact["email"], self.contact["email"])
        self.assertIn("experience", parsed.sections)
        critical = [i.title for i in parsed.issues if i.severity == "critical"]
        self.assertEqual(critical, [])

    def test_tailored_text_render(self):
        from resume.tailor import rendered_text

        text = rendered_text(self.tailored, self.contact)
        self.assertIn("Ayesha Rahman", text)
        self.assertIn("Kubernetes", text)
        self.assertIn("EXPERIENCE", text)


@override_settings(MEDIA_ROOT=MEDIA, GROQ_API_KEY="", LLM_BACKEND="echo")
class ApiTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user("api", password="pw12345678")
        self.client.login(username="api", password="pw12345678")

    def test_document_list_requires_login(self):
        self.client.logout()
        self.assertIn(self.client.get("/api/documents/").status_code, (401, 403))

    def test_resumable_upload_round_trip(self):
        data = build_test_pdf()
        half = len(data) // 2
        parts = [data[:half], data[half:]]

        init = self.client.post(
            "/api/documents/upload/init/",
            {"filename": "big.pdf", "total_size": len(data), "total_parts": 2},
        )
        self.assertEqual(init.status_code, 201, init.content)
        upload_id = init.json()["upload_id"]

        # a part sent twice must not corrupt the assembly
        for index, blob in enumerate(parts):
            for _ in range(2 if index == 0 else 1):
                resp = self.client.post(
                    f"/api/documents/upload/{upload_id}/part/",
                    {"index": index, "file": ContentFile(blob, name=f"p{index}")},
                )
                self.assertEqual(resp.status_code, 200, resp.content)

        status = self.client.get(f"/api/documents/upload/{upload_id}/status/").json()
        self.assertEqual(status["missing"], [])

        done = self.client.post(f"/api/documents/upload/{upload_id}/complete/")
        self.assertEqual(done.status_code, 201, done.content)

        doc = Document.objects.get(id=done.json()["id"])
        self.assertEqual(doc.status, Document.Status.READY, doc.error)
        self.assertEqual(doc.page_count, 3)

    def test_incomplete_upload_reports_missing_parts(self):
        init = self.client.post(
            "/api/documents/upload/init/",
            {"filename": "x.pdf", "total_size": 500, "total_parts": 3},
        ).json()
        resp = self.client.post(f"/api/documents/upload/{init['upload_id']}/complete/")
        self.assertEqual(resp.status_code, 409)
        self.assertEqual(resp.json()["missing"], [0, 1, 2])

    def test_oversized_upload_is_rejected_with_a_useful_message(self):
        resp = self.client.post(
            "/api/documents/upload/init/",
            {"filename": "huge.pdf", "total_size": 10 ** 12, "total_parts": 5},
        )
        self.assertEqual(resp.status_code, 400)
        self.assertIn("MB", str(resp.json()))

    def test_figures_endpoint(self):
        doc = Document(owner=self.user, title="r.pdf")
        doc.file.save("r.pdf", ContentFile(build_test_pdf()), save=True)
        process_document(doc.id)
        resp = self.client.get(f"/api/documents/{doc.id}/figures/")
        self.assertEqual(resp.status_code, 200)
        self.assertIn("figures", resp.json())

    def test_resume_upload_and_score(self):
        resp = self.client.post(
            "/api/resume/",
            {"title": "My CV", "file": ContentFile(CV_TEXT.encode(), name="cv.txt")},
        )
        self.assertEqual(resp.status_code, 201, resp.content)
        body = resp.json()
        self.assertEqual(body["status"], "ready")
        self.assertGreater(body["base_score"]["overall"], 0)
        self.assertEqual(body["contact"]["email"], "ayesha.rahman@example.com")

    def test_tailoring_without_a_key_still_returns_score_and_gaps(self):
        """The audit is local. Only the rewriting needs a model."""
        cv = self.client.post(
            "/api/resume/",
            {"title": "CV", "file": ContentFile(CV_TEXT.encode(), name="cv.txt")},
        ).json()

        resp = self.client.post(
            "/api/resume/tailorings/",
            {"resume": cv["id"], "job_title": "Senior SRE",
             "job_description": JOB_DESCRIPTION},
        )
        self.assertEqual(resp.status_code, 201, resp.content)
        body = resp.json()
        self.assertEqual(body["status"], "ready")
        self.assertGreater(body["before_score"]["overall"], 0)
        self.assertTrue(body["gap_analysis"]["missing_keywords"])
        self.assertTrue(body["requirements"]["must_have"])

    def test_short_job_description_is_rejected(self):
        cv = self.client.post(
            "/api/resume/",
            {"title": "CV", "file": ContentFile(CV_TEXT.encode(), name="cv.txt")},
        ).json()
        resp = self.client.post(
            "/api/resume/tailorings/",
            {"resume": cv["id"], "job_description": "SRE wanted"},
        )
        self.assertEqual(resp.status_code, 400)

    def test_cannot_tailor_someone_elses_cv(self):
        User.objects.create_user("other", password="pw12345678")
        cv = self.client.post(
            "/api/resume/",
            {"title": "CV", "file": ContentFile(CV_TEXT.encode(), name="cv.txt")},
        ).json()
        self.client.logout()
        self.client.login(username="other", password="pw12345678")
        resp = self.client.post(
            "/api/resume/tailorings/",
            {"resume": cv["id"], "job_description": JOB_DESCRIPTION},
        )
        self.assertEqual(resp.status_code, 403)


class BoilerplateKeywordTests(TestCase):
    """Application-process wording must not be scored as job requirements."""

    JD = (
        "Vacancy Announcement\n"
        "Title: Supply Chain Associate\n"
        "Minimum 3 years of experience in procurement and warehouse operations\n"
        "Hands-on knowledge of ERP systems such as SAP or Oracle\n"
        "UNHCR is an equal opportunity employer.\n"
        "Qualified female candidates are strongly encouraged to apply.\n"
        "Only shortlisted candidates will be contacted for interview.\n"
        "Please note applications submitted after the closing date will not be considered.\n"
        "Kindly apply through the link provided below.\n"
    )

    def test_process_words_are_not_keywords(self):
        from resume.ats import extract_requirements

        kws = set(extract_requirements(self.JD)["keywords"])
        for term in [
            "apply", "shortlisted", "candidates", "contacted", "interview",
            "please", "kindly", "note", "below", "submitted", "considered",
            "gender", "female", "diversity", "vacancy", "only", "provided",
        ]:
            self.assertNotIn(term, kws, f"{term!r} is boilerplate, not a requirement")

    def test_real_requirements_survive(self):
        from resume.ats import extract_requirements

        kws = set(extract_requirements(self.JD)["keywords"])
        for term in ["procurement", "warehouse", "supply", "chain", "oracle"]:
            self.assertIn(term, kws, f"{term!r} is a real requirement and was dropped")

    def test_boilerplate_lines_removed_before_parsing(self):
        from resume.ats import strip_boilerplate

        body = strip_boilerplate(self.JD)
        self.assertNotIn("shortlisted", body.lower())
        self.assertIn("procurement", body.lower())


class SynonymTests(TestCase):
    def test_equivalent_wording_counts_as_evidence(self):
        from resume import synonyms

        cv = {"purchasing", "stock control", "led", "dashboards"}
        for term in ["procurement", "inventory", "manage", "reporting"]:
            self.assertTrue(synonyms.matches(term, cv), f"{term} should match {cv}")

    def test_unrelated_skills_never_match(self):
        from resume import synonyms

        cv = {"python", "django"}
        for term in ["ruby", "kubernetes", "welding"]:
            self.assertFalse(synonyms.matches(term, cv))


class FabricationTests(TestCase):
    SOURCE = (
        "Warehouse Officer, Beta Traders, 2018 - 2020\n"
        "- Ran daily cycle counts across three zones\n"
        "Skills: Excel, stock control\n"
    )

    def _cv(self, **over):
        base = {
            "summary": "Warehouse officer with stock control experience.",
            "skills": ["Excel", "Stock control"],
            "experience": [{
                "company": "Beta Traders", "title": "Warehouse Officer",
                "dates": "2018 - 2020",
                "bullets": ["Ran daily cycle counts across three zones."],
            }],
            "projects": [], "education": [], "certifications": [],
        }
        base.update(over)
        return base

    def test_honest_rewrite_is_clean(self):
        from resume import verify

        self.assertTrue(verify.check(self._cv(), self.SOURCE)["clean"])

    def test_invented_employer_is_critical(self):
        from resume import verify

        cv = self._cv(experience=[{
            "company": "Acme Logistics", "title": "Warehouse Officer",
            "dates": "2018 - 2020", "bullets": ["Ran cycle counts."],
        }])
        rep = verify.check(cv, self.SOURCE)
        self.assertFalse(rep["clean"])
        self.assertEqual(rep["counts"]["critical"], 1)

    def test_invented_skill_is_flagged(self):
        from resume import verify

        rep = verify.check(self._cv(skills=["Excel", "Kubernetes"]), self.SOURCE)
        self.assertTrue(any(f["value"] == "Kubernetes" for f in rep["findings"]))

    def test_invented_number_is_flagged(self):
        from resume import verify

        cv = self._cv(experience=[{
            "company": "Beta Traders", "title": "Warehouse Officer",
            "dates": "2018 - 2020",
            "bullets": ["Ran cycle counts across 4,000 SKUs."],
        }])
        rep = verify.check(cv, self.SOURCE)
        self.assertTrue(any(f["kind"] == "number" for f in rep["findings"]))

    def test_user_supplied_answers_clear_the_flag(self):
        """A fact the user asserted in the interview is not a fabrication."""
        from resume import verify

        cv = self._cv(experience=[{
            "company": "Beta Traders", "title": "Warehouse Officer",
            "dates": "2018 - 2020",
            "bullets": ["Ran cycle counts across 4,000 SKUs."],
        }])
        rep = verify.check(cv, self.SOURCE, extra_sources=["I counted 4,000 SKUs weekly"])
        self.assertFalse(any(f["kind"] == "number" for f in rep["findings"]))


class InterviewTests(TestCase):
    GAPS = {"unmet_requirements": [
        {"requirement": "Minimum 3 years of relevant experience in procurement", "coverage": 0},
        {"requirement": "Hands-on knowledge of ERP systems such as SAP", "coverage": 10},
    ]}

    def test_questions_are_generated_without_a_key(self):
        from resume import interview

        qs = interview.build_questions({}, self.GAPS)
        self.assertEqual(len(qs), 2)
        for q in qs:
            self.assertTrue(q["question"].endswith("?"))

    def test_recruiter_preamble_is_stripped(self):
        from resume import interview

        qs = interview.build_questions({}, self.GAPS)
        self.assertNotIn("minimum 3 years", qs[0]["question"].lower())
        self.assertIn("procurement", qs[0]["question"].lower())

    def test_acronyms_keep_their_case(self):
        from resume import interview

        qs = interview.build_questions({}, self.GAPS)
        self.assertIn("ERP", qs[1]["question"])

    def test_negative_answers_never_become_evidence(self):
        from resume import interview

        qs = interview.build_questions({}, self.GAPS)
        ev = interview.answers_to_evidence(qs, {"0": "no", "1": "  "})
        self.assertEqual(ev, [])

    def test_real_answers_become_evidence(self):
        from resume import interview

        qs = interview.build_questions({}, self.GAPS)
        ev = interview.answers_to_evidence(qs, {"0": "Six years buying for two warehouses"})
        self.assertEqual(len(ev), 1)
        self.assertIn("Six years", interview.evidence_text(ev))


class PdfExportTests(TestCase):
    CV = {
        "summary": "Supply chain associate with six years in procurement across two sites.",
        "skills": ["Procurement", "Inventory management", "SAP MM"],
        "experience": [{
            "company": "Acme Logistics", "title": "Supply Chain Associate",
            "dates": "2020 - 2026",
            "bullets": [
                "Managed procurement for 4,000 SKUs across two warehouses, from requisition to payment.",
                "Reconciled stock monthly, cutting variance from 6% to under 1%.",
            ],
        }],
        "projects": [], "education": ["BSc Supply Chain, 2018"], "certifications": [],
    }
    CONTACT = {"name": "Test Candidate", "email": "t@example.com", "phone": "0300"}

    def test_pdf_text_is_extractable_and_complete(self):
        from resume import export

        pdf = export.build_pdf(self.CV, self.CONTACT)
        report = export.verify_pdf_ats_safe(pdf)
        self.assertTrue(report["safe"], report["problems"])

        import fitz
        text = "".join(page.get_text() for page in fitz.open(stream=pdf, filetype="pdf"))
        for probe in ["Test Candidate", "4,000 SKUs", "Reconciled", "BSc Supply Chain"]:
            self.assertIn(probe, text, f"{probe!r} missing from the PDF")

    def test_pdf_carries_no_images(self):
        from resume import export

        report = export.verify_pdf_ats_safe(export.build_pdf(self.CV, self.CONTACT))
        self.assertNotIn("contains images", report["problems"])

    def test_long_bullets_are_not_truncated(self):
        """A bullet longer than the estimated box must still render in full."""
        from resume import export

        long_bullet = (
            "Coordinated end-to-end customs clearance and freight forwarding for inbound "
            "shipments from six countries, resolving documentation discrepancies with "
            "brokers and cutting average dwell time at port by four days across the year."
        )
        cv = dict(self.CV)
        cv["experience"] = [{
            "company": "Acme", "title": "Officer", "dates": "2020",
            "bullets": [long_bullet],
        }]
        pdf = export.build_pdf(cv, self.CONTACT)

        import fitz
        text = "".join(p.get_text() for p in fitz.open(stream=pdf, filetype="pdf"))
        flat = " ".join(text.split())
        self.assertIn("cutting average dwell time at port by four days", flat)


class InterviewRescoreTests(TestCase):
    """Answers must be worth something even when no model is available."""

    CV = (
        "Ahsaan Munir\nahsaan@example.com | 0300 1234567\n\n"
        "EXPERIENCE\nWarehouse Officer, Beta Traders, 2018 - 2020\n"
        "- Ran daily cycle counts across three zones\n"
        "- Handled purchasing for consumable stock\n\n"
        "EDUCATION\nBSc Supply Chain Management, 2018\n\n"
        "SKILLS\nExcel, stock control\n"
    )
    JD = (
        "Supply Chain Associate\n"
        "- Minimum 3 years of relevant experience in procurement and warehouse operations\n"
        "- Proven experience with inventory management and stock reconciliation\n"
        "- Hands-on knowledge of ERP systems, preferably SAP or Oracle\n"
        "- Demonstrated ability in vendor management and tender evaluation\n"
        "Only shortlisted candidates will be contacted.\n"
    )

    def _run(self, answers=None):
        from django.contrib.auth.models import User
        from django.core.files.base import ContentFile
        from resume import services
        from resume.models import Resume, Tailoring

        user = User.objects.create_user(f"u{id(answers)}", password="x")
        resume = Resume.objects.create(owner=user, title="CV")
        resume.file.save("cv.txt", ContentFile(self.CV.encode()), save=True)
        services.process_resume(resume.id)

        job = Tailoring.objects.create(
            owner=user, resume=resume, job_description=self.JD, answers=answers or {}
        )
        services.process_tailoring(job.id)
        job.refresh_from_db()
        return job

    def test_questions_appear_without_a_key(self):
        job = self._run()
        self.assertEqual(job.status, "ready")
        self.assertTrue(job.questions, "no interview questions were generated")

    def test_relevant_answer_raises_the_score(self):
        base = self._run()
        before = base.before_score["overall"]

        answered = self._run(answers={
            "0": "I used SAP MM for four years raising purchase orders and running "
                 "tender evaluation for vendor selection.",
        })
        self.assertTrue(answered.evidence, "the answer was discarded")
        after = answered.after_score.get("overall")
        self.assertIsNotNone(after)
        self.assertGreater(after, before)

    def test_synonyms_close_gaps_without_rewording(self):
        """'purchasing' in the CV should satisfy 'procurement' in the posting."""
        job = self._run()
        matched = job.gap_analysis["matched_keywords"]
        self.assertIn("procurement", matched)
