"""
CV parsing and ATS format auditing.

Most "AI CV" tools only rewrite wording. That misses the failure that actually
loses interviews: an applicant tracking system parsing the file wrongly before
a human ever reads it. A two-column CV commonly arrives at the recruiter as
interleaved nonsense; a job title inside a text box or a header often does not
arrive at all.

So this module does two separate jobs:

  parse()        pull structure out - contact, sections, skills, bullets
  audit_format() find the things that break machine parsing, with evidence

The audit is deterministic. No model is asked whether a layout has two
columns; that is measured from the geometry of the text on the page.
"""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

try:
    import pymupdf
except ImportError:  # pragma: no cover
    import fitz as pymupdf


# ------------------------------------------------------------------ patterns

EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]{2,}")
PHONE = re.compile(r"(?:\+?\d[\d\s\-().]{7,}\d)")
URL = re.compile(r"(?:https?://|www\.)[^\s<>,;]+", re.I)
LINKEDIN = re.compile(r"(?:linkedin\.com/in/|linkedin:\s*)([\w\-]+)", re.I)
GITHUB = re.compile(r"(?:github\.com/)([\w\-]+)", re.I)

DATE_RANGE = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}"
    r"|\d{1,2}/\d{4}|\d{4})"
    r"\s*(?:-|–|—|to|until)\s*"
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}"
    r"|\d{1,2}/\d{4}|\d{4}|present|current|now|ongoing)",
    re.I,
)

BULLET_MARK = re.compile(r"^\s*[\u2022\u25cf\u25aa\u2023\u2043\u00b7•▪◦‣●○*\-–—]\s+")

# Headings ATS parsers reliably recognise, and the synonyms people use instead.
CANONICAL_SECTIONS = {
    "summary": ["summary", "professional summary", "profile", "about me", "objective",
                "career objective", "personal statement", "executive summary"],
    "experience": ["experience", "work experience", "professional experience",
                   "employment", "employment history", "career history",
                   "work history", "relevant experience"],
    "education": ["education", "academic background", "qualifications",
                  "academic qualifications", "educational background"],
    "skills": ["skills", "technical skills", "core competencies", "competencies",
               "key skills", "areas of expertise", "technologies", "tech stack"],
    "projects": ["projects", "key projects", "selected projects", "personal projects",
                 "portfolio"],
    "certifications": ["certifications", "certificates", "licenses",
                       "certifications and licenses", "professional development"],
    "awards": ["awards", "honors", "honours", "achievements", "recognition"],
    "publications": ["publications", "papers", "research"],
    "languages": ["languages", "language skills"],
    "volunteer": ["volunteer", "volunteering", "community involvement"],
    "interests": ["interests", "hobbies", "activities"],
    "references": ["references", "referees"],
}

SECTION_LOOKUP = {
    alias: canon for canon, aliases in CANONICAL_SECTIONS.items() for alias in aliases
}

ACTION_VERBS = {
    "achieved", "administered", "advised", "analysed", "analyzed", "architected",
    "authored", "automated", "built", "championed", "coached", "conducted",
    "consolidated", "coordinated", "created", "cut", "delivered", "deployed",
    "designed", "developed", "directed", "drove", "eliminated", "engineered",
    "established", "executed", "expanded", "facilitated", "founded", "generated",
    "grew", "implemented", "improved", "increased", "initiated", "integrated",
    "introduced", "launched", "led", "maintained", "managed", "mentored",
    "migrated", "modernised", "modernized", "negotiated", "optimised",
    "optimized", "orchestrated", "overhauled", "owned", "partnered", "pioneered",
    "planned", "produced", "programmed", "raised", "rebuilt", "reduced",
    "refactored", "resolved", "restructured", "scaled", "secured", "shipped",
    "simplified", "spearheaded", "standardised", "standardized", "streamlined",
    "supervised", "supported", "trained", "transformed", "upgraded", "wrote",
}

WEAK_OPENERS = {
    "responsible", "worked", "helped", "assisted", "involved", "participated",
    "tasked", "duties", "handled", "dealt", "various", "etc",
}

QUANTIFIED = re.compile(r"(\d+(?:\.\d+)?\s*%|\$\s?\d|\b\d{2,}\b|\b\d+(?:\.\d+)?\s*(?:x|k|m|bn|million|billion|hours?|days?|weeks?|months?|users?|customers?|clients?|people|engineers?)\b)", re.I)

RISKY_GLYPHS = re.compile(r"[\u2500-\u257F\u2580-\u259F\u25A0-\u25FF\uE000-\uF8FF\U0001F300-\U0001FAFF]")

SAFE_FONTS = {
    "arial", "calibri", "helvetica", "times new roman", "times", "georgia",
    "garamond", "cambria", "verdana", "tahoma", "trebuchet ms", "book antiqua",
    "palatino", "lato", "roboto", "open sans", "liberation serif",
    "liberation sans", "carlito", "dejavu sans", "dejavu serif",
}


@dataclass
class FormatIssue:
    code: str
    severity: str           # "critical" | "warning" | "note"
    title: str
    detail: str
    fix: str
    evidence: str = ""


@dataclass
class ParsedResume:
    raw_text: str = ""
    contact: Dict = field(default_factory=dict)
    sections: Dict[str, List[str]] = field(default_factory=dict)
    section_order: List[str] = field(default_factory=list)
    skills: List[str] = field(default_factory=list)
    bullets: List[str] = field(default_factory=list)
    page_count: int = 0
    word_count: int = 0
    issues: List[FormatIssue] = field(default_factory=list)
    stats: Dict = field(default_factory=dict)


# ------------------------------------------------------------ text extraction

def _pdf_text_and_geometry(path: str) -> Tuple[str, Dict]:
    doc = pymupdf.open(path)
    try:
        pages_text: List[str] = []
        fonts: Counter = Counter()
        image_count = 0
        vector_paths = 0
        column_pages = 0
        header_lines: Counter = Counter()
        footer_lines: Counter = Counter()
        tiny_text = 0
        total_spans = 0

        for i in range(doc.page_count):
            page = doc.load_page(i)
            pages_text.append(page.get_text("text") or "")

            try:
                image_count += len(page.get_images(full=True))
                vector_paths += len(page.get_drawings())
            except Exception:
                pass

            data = page.get_text("dict")
            xs: List[Tuple[float, float]] = []
            height = page.rect.height
            for block in data.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    lx0, ly0, lx1, ly1 = line.get("bbox", (0, 0, 0, 0))
                    line_text = "".join(
                        span.get("text", "") for span in line.get("spans", [])
                    ).strip()
                    if line_text:
                        # measure per LINE, not per block: a CV laid out with two
                        # large text frames yields only two blocks, far too few
                        # to see a gutter, but dozens of lines
                        xs.append((lx0, lx1))
                        if ly1 < height * 0.08:
                            header_lines[line_text[:60]] += 1
                        if ly0 > height * 0.93:
                            footer_lines[line_text[:60]] += 1
                    for span in line.get("spans", []):
                        total_spans += 1
                        fonts[(span.get("font") or "").split("+")[-1].split("-")[0].lower()] += 1
                        if span.get("size", 12) < 8.0:
                            tiny_text += 1

            if _looks_multi_column(xs, page.rect.width):
                column_pages += 1

        geometry = {
            "fonts": fonts,
            "image_count": image_count,
            "vector_paths": vector_paths,
            "column_pages": column_pages,
            "page_count": doc.page_count,
            "repeated_headers": [t for t, n in header_lines.items() if n >= max(2, doc.page_count - 1)],
            "repeated_footers": [t for t, n in footer_lines.items() if n >= max(2, doc.page_count - 1)],
            "tiny_text_spans": tiny_text,
            "total_spans": total_spans,
        }
        return "\n".join(pages_text), geometry
    finally:
        doc.close()


def _looks_multi_column(xs: List[Tuple[float, float]], page_width: float) -> bool:
    """Detect side-by-side columns from block x-ranges.

    A single-column CV has blocks starting at roughly one left margin. A
    two-column CV has two clusters of left edges with a gutter between them,
    and many ATS parsers read straight across that gutter, splicing a job
    title onto an unrelated skill.
    """
    if len(xs) < 10:
        return False
    mid = page_width / 2
    # a line belongs to the left column if it ENDS before the midline, and to
    # the right column if it STARTS after it; anything spanning the middle is
    # full-width and votes against a column layout
    left = [x for x in xs if x[1] < mid * 1.02]
    right = [x for x in xs if x[0] > mid * 0.98]
    spanning = [x for x in xs if x[0] < mid * 0.98 and x[1] > mid * 1.02]

    if len(left) < 4 or len(right) < 4:
        return False
    # a handful of right-aligned dates is not a column
    if len(right) < len(xs) * 0.12:
        return False
    return len(spanning) < len(xs) * 0.35


def _docx_text_and_geometry(path: str) -> Tuple[str, Dict]:
    from docx import Document as Docx

    d = Docx(path)
    parts: List[str] = [p.text for p in d.paragraphs]

    table_cells = 0
    for table in d.tables:
        for row in table.rows:
            table_cells += len(row.cells)
            parts.append(" | ".join(c.text.strip() for c in row.cells))

    xml = d.element.xml
    fonts: Counter = Counter()
    for run in d.element.iter():
        if run.tag.endswith("}rFonts"):
            for key in ("ascii", "hAnsi"):
                val = run.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + key
                )
                if val:
                    fonts[val.lower()] += 1

    header_text: List[str] = []
    footer_text: List[str] = []
    for section in d.sections:
        try:
            header_text += [p.text for p in section.header.paragraphs if p.text.strip()]
            footer_text += [p.text for p in section.footer.paragraphs if p.text.strip()]
        except Exception:
            pass

    geometry = {
        "fonts": fonts,
        "image_count": xml.count("<w:drawing") + xml.count("<w:pict"),
        "vector_paths": 0,
        "column_pages": 1 if "<w:cols" in xml and 'w:num="2"' in xml else 0,
        "page_count": 0,
        "table_cells": table_cells,
        "text_boxes": xml.count("txbxContent"),
        "repeated_headers": header_text,
        "repeated_footers": footer_text,
        "tiny_text_spans": 0,
        "total_spans": 0,
    }
    return "\n".join(parts), geometry


def extract(path: str) -> Tuple[str, Dict]:
    ext = path.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return _pdf_text_and_geometry(path)
    if ext in ("docx", "doc"):
        return _docx_text_and_geometry(path)
    with open(path, "r", encoding="utf-8", errors="ignore") as fh:
        return fh.read(), {"fonts": Counter(), "image_count": 0, "vector_paths": 0,
                           "column_pages": 0, "page_count": 0, "repeated_headers": [],
                           "repeated_footers": [], "tiny_text_spans": 0, "total_spans": 0}


# ------------------------------------------------------------------ parsing

def _find_contact(text: str) -> Dict:
    head = text[:1400]
    emails = EMAIL.findall(text)
    phones = [p.strip() for p in PHONE.findall(text) if len(re.sub(r"\D", "", p)) >= 9]
    li = LINKEDIN.search(text)
    gh = GITHUB.search(text)

    name = ""
    for line in head.splitlines():
        line = line.strip()
        if not line or EMAIL.search(line) or PHONE.search(line) or URL.search(line):
            continue
        words = line.split()
        if 1 < len(words) <= 5 and all(w[:1].isalpha() for w in words) and len(line) < 60:
            if not SECTION_LOOKUP.get(line.lower().strip(":")):
                name = line
                break

    return {
        "name": name,
        "email": emails[0] if emails else "",
        "phone": phones[0] if phones else "",
        "linkedin": f"linkedin.com/in/{li.group(1)}" if li else "",
        "github": f"github.com/{gh.group(1)}" if gh else "",
        "location": _guess_location(head),
    }


def _guess_location(head: str) -> str:
    for line in head.splitlines()[:8]:
        line = line.strip()
        if 3 < len(line) < 60 and "," in line and not EMAIL.search(line) and not URL.search(line):
            if re.search(r"[A-Z][a-z]+,\s*[A-Z]", line):
                return line
    return ""


def _is_heading(line: str) -> Optional[str]:
    clean = line.strip().strip(":").strip()
    if not clean or len(clean) > 48:
        return None
    key = re.sub(r"[^a-z ]", "", clean.lower()).strip()
    if key in SECTION_LOOKUP:
        return SECTION_LOOKUP[key]
    # ALL CAPS short line is a heading even when worded unusually
    if clean.isupper() and 3 <= len(clean) <= 40:
        return SECTION_LOOKUP.get(key, "other:" + clean.lower())
    return None


def _split_sections(text: str) -> Tuple[Dict[str, List[str]], List[str]]:
    sections: Dict[str, List[str]] = {}
    order: List[str] = []
    current = "header"
    sections[current] = []
    order.append(current)

    for raw in text.splitlines():
        heading = _is_heading(raw)
        if heading:
            current = heading
            sections.setdefault(current, [])
            if current not in order:
                order.append(current)
            continue
        if raw.strip():
            sections.setdefault(current, []).append(raw.rstrip())
    return sections, order


def _extract_skills(sections: Dict[str, List[str]]) -> List[str]:
    lines = sections.get("skills", [])
    blob = " ".join(lines)
    parts = re.split(r"[,;|•·\u2022\n]+|\s{3,}", blob)
    skills = []
    for p in parts:
        p = p.strip(" .:-–—")
        p = re.sub(r"^(?:proficient in|experienced with|familiar with)\s+", "", p, flags=re.I)
        if 1 < len(p) <= 42 and not p.lower().startswith(("years", "including")):
            skills.append(p)
    seen = set()
    out = []
    for s in skills:
        k = s.lower()
        if k not in seen:
            seen.add(k)
            out.append(s)
    return out[:80]


def _collect_bullets(sections: Dict[str, List[str]]) -> List[str]:
    bullets = []
    for key in ("experience", "projects", "volunteer", "other"):
        for line in sections.get(key, []):
            if BULLET_MARK.match(line):
                bullets.append(BULLET_MARK.sub("", line).strip())
            elif len(line) > 60 and line[:1].isupper() and not DATE_RANGE.search(line):
                bullets.append(line.strip())
    return bullets


# -------------------------------------------------------------- format audit

def audit_format(text: str, geometry: Dict, filename: str) -> List[FormatIssue]:
    issues: List[FormatIssue] = []
    ext = filename.lower().rsplit(".", 1)[-1]

    if geometry.get("column_pages"):
        issues.append(FormatIssue(
            code="multi_column",
            severity="critical",
            title="Multi-column layout detected",
            detail=(f"{geometry['column_pages']} page(s) place text in side-by-side "
                    "columns. Many applicant tracking systems read straight across the "
                    "gutter, splicing a job title onto an unrelated line."),
            fix="Rebuild as a single column running top to bottom.",
            evidence=f"{geometry['column_pages']} page(s) affected",
        ))

    if geometry.get("table_cells", 0) > 4:
        issues.append(FormatIssue(
            code="tables",
            severity="critical",
            title="Content inside tables",
            detail=(f"About {geometry['table_cells']} table cells hold CV content. "
                    "Table parsing is the single most common cause of scrambled "
                    "output from an ATS."),
            fix="Move the content into ordinary paragraphs and bullet lists.",
        ))

    if geometry.get("text_boxes", 0):
        issues.append(FormatIssue(
            code="text_boxes",
            severity="critical",
            title="Text boxes in the document",
            detail=("Text inside a text box is often skipped completely. If a job "
                    "title sits in one, the parser records no job title."),
            fix="Delete the text boxes and place the text in the body.",
        ))

    if geometry.get("image_count", 0) > 0:
        issues.append(FormatIssue(
            code="images",
            severity="warning",
            title=f"{geometry['image_count']} image(s) embedded",
            detail=("Text inside an image is invisible to a parser. A photo also "
                    "invites bias screening in the UK, US and much of Europe."),
            fix="Remove photos, logos and icons; keep information as real text.",
        ))

    if geometry.get("repeated_headers") or geometry.get("repeated_footers"):
        sample = (geometry.get("repeated_headers") or geometry.get("repeated_footers"))[:1]
        issues.append(FormatIssue(
            code="header_footer",
            severity="warning",
            title="Contact details or content in a header/footer",
            detail=("Headers and footers are frequently dropped. If your phone number "
                    "or email lives there, the recruiter may get a CV with no way to "
                    "contact you."),
            fix="Move everything into the main body, at the top of page one.",
            evidence=str(sample[0])[:80] if sample else "",
        ))

    fonts = geometry.get("fonts") or Counter()
    odd = [f for f in fonts if f and not any(s in f for s in SAFE_FONTS)]
    if odd and sum(fonts[f] for f in odd) > sum(fonts.values()) * 0.3:
        issues.append(FormatIssue(
            code="fonts",
            severity="note",
            title="Unusual fonts",
            detail=f"Most text uses {', '.join(sorted(set(odd))[:3])}. Uncommon or "
                   "decorative fonts occasionally extract as wrong characters.",
            fix="Use Arial, Calibri, Georgia or Times New Roman.",
        ))

    if geometry.get("tiny_text_spans", 0) > max(8, geometry.get("total_spans", 0) * 0.08):
        issues.append(FormatIssue(
            code="tiny_text",
            severity="warning",
            title="Very small text",
            detail="A noticeable amount of text is under 8pt, which suggests content "
                   "squeezed to fit. It reads badly on screen and prints worse.",
            fix="Set body text to 10-12pt and cut content instead of shrinking it.",
        ))

    glyphs = RISKY_GLYPHS.findall(text)
    if len(glyphs) > 6:
        issues.append(FormatIssue(
            code="glyphs",
            severity="note",
            title="Decorative symbols and icons",
            detail=f"{len(glyphs)} box-drawing, icon-font or emoji characters found. "
                   "These often extract as question marks.",
            fix="Replace icons with plain words, and use a standard bullet.",
        ))

    if ext not in ("pdf", "docx"):
        issues.append(FormatIssue(
            code="filetype",
            severity="warning",
            title=f".{ext} is an unusual CV format",
            detail="Most systems accept PDF and DOCX. Anything else risks rejection "
                   "at the upload step.",
            fix="Submit a DOCX unless the posting asks for PDF.",
        ))

    if geometry.get("page_count", 0) > 3:
        issues.append(FormatIssue(
            code="length",
            severity="note",
            title=f"{geometry['page_count']} pages",
            detail="Beyond two pages, later content is read less often - by people "
                   "and by keyword-ranking alike.",
            fix="Cut to two pages unless you are in academia or a senior role with "
                "a publication record.",
        ))

    return issues


# ---------------------------------------------------------------- entry point

def parse(path: str, filename: str = "") -> ParsedResume:
    filename = filename or path
    text, geometry = extract(path)
    sections, order = _split_sections(text)
    bullets = _collect_bullets(sections)

    words = len(text.split())
    page_count = geometry.get("page_count") or max(1, words // 500)

    quantified = sum(1 for b in bullets if QUANTIFIED.search(b))
    strong = sum(1 for b in bullets if b.split()[0].lower().rstrip(":,.") in ACTION_VERBS) if bullets else 0
    weak = sum(1 for b in bullets if b.split()[0].lower().rstrip(":,.") in WEAK_OPENERS) if bullets else 0

    return ParsedResume(
        raw_text=text,
        contact=_find_contact(text),
        sections=sections,
        section_order=order,
        skills=_extract_skills(sections),
        bullets=bullets,
        page_count=page_count,
        word_count=words,
        issues=audit_format(text, geometry, filename),
        stats={
            "bullets": len(bullets),
            "quantified_bullets": quantified,
            "strong_openers": strong,
            "weak_openers": weak,
            "dates_found": len(DATE_RANGE.findall(text)),
            "sections_found": [s for s in order if not s.startswith("other:")],
        },
    )
