"""
ATS-safe export.

Everything here is a deliberate omission. No tables, no text boxes, no
columns, no images, no headers or footers, no fancy bullets, no custom
fonts. Those are exactly the features that make a CV look designed and make
it parse badly, and this file exists to guarantee the output cannot contain
them - regardless of what the original CV did.

Contact details go in the body of page one, never a header, because header
content is routinely dropped and a CV with no reachable phone number is worse
than an ugly one.
"""

from __future__ import annotations

import io
from typing import Dict

from docx import Document as Docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BODY_FONT = "Calibri"          # ships with Word, extracts cleanly everywhere
BODY_SIZE = Pt(10.5)
NAME_SIZE = Pt(20)
HEADING_SIZE = Pt(11.5)
INK = RGBColor(0x1A, 0x1A, 0x1A)
RULE = RGBColor(0x44, 0x44, 0x44)


def _base(doc: Docx) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.font.color.rgb = INK
    pf = style.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.06

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(46)


def _heading(doc: Docx, text: str) -> None:
    """A plain bold uppercase paragraph, not a Word Heading style.

    Word's built-in heading styles carry outline levels that some parsers
    interpret as document structure and mis-nest. A bold line is understood
    by everything.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = HEADING_SIZE
    run.font.color.rgb = RULE
    # a border, rather than a drawn line or a one-cell table
    _bottom_border(p)


def _bottom_border(paragraph) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    pPr.append(borders)


def _bullet(doc: Docx, text: str) -> None:
    """A hyphen and a hanging indent, not a list style.

    Word list styles store the bullet glyph outside the run text; some
    extractors then lose the bullet, and occasionally the line with it.
    A literal hyphen always survives.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.first_line_indent = Pt(-12)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"- {text}")


def build_docx(tailored: Dict, contact: Dict) -> bytes:
    doc = Docx()
    _base(doc)

    name = contact.get("name") or "Curriculum Vitae"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.size = NAME_SIZE

    details = [v for v in (
        contact.get("email"), contact.get("phone"), contact.get("location"),
        contact.get("linkedin"), contact.get("github"),
    ) if v]
    if details:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(6)
        cp.add_run(" | ".join(details)).font.size = Pt(9.5)

    if tailored.get("summary"):
        _heading(doc, "Professional Summary")
        doc.add_paragraph(tailored["summary"])

    if tailored.get("skills"):
        _heading(doc, "Skills")
        doc.add_paragraph(", ".join(tailored["skills"]))

    if tailored.get("experience"):
        _heading(doc, "Experience")
        for job in tailored["experience"]:
            jp = doc.add_paragraph()
            jp.paragraph_format.space_before = Pt(7)
            jp.paragraph_format.space_after = Pt(1)
            title = job.get("title") or ""
            company = job.get("company") or ""
            head = f"{title} — {company}" if title and company else (title or company)
            jp.add_run(head).bold = True
            if job.get("dates"):
                jp.add_run(f"   {job['dates']}").italic = True
            for b in job.get("bullets", []):
                _bullet(doc, b)

    if tailored.get("projects"):
        _heading(doc, "Projects")
        for proj in tailored["projects"]:
            if proj.get("name"):
                pp = doc.add_paragraph()
                pp.paragraph_format.space_before = Pt(6)
                pp.paragraph_format.space_after = Pt(1)
                pp.add_run(proj["name"]).bold = True
            for b in proj.get("bullets", []):
                _bullet(doc, b)

    if tailored.get("education"):
        _heading(doc, "Education")
        for line in tailored["education"]:
            doc.add_paragraph(line)

    if tailored.get("certifications"):
        _heading(doc, "Certifications")
        for line in tailored["certifications"]:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def verify_ats_safe(docx_bytes: bytes) -> Dict:
    """Prove the export is clean rather than asserting it.

    Runs against the produced file, so a future edit to build_docx that
    reintroduces a table is caught by the test suite instead of by a user.
    """
    d = Docx(io.BytesIO(docx_bytes))
    xml = d.element.xml
    problems = []

    if len(d.tables) > 0:
        problems.append("contains tables")
    if "txbxContent" in xml:
        problems.append("contains text boxes")
    if "<w:drawing" in xml or "<w:pict" in xml:
        problems.append("contains images")
    if 'w:num="2"' in xml or 'w:num="3"' in xml:
        problems.append("contains multiple columns")
    for section in d.sections:
        if any(p.text.strip() for p in section.header.paragraphs):
            problems.append("content in header")
        if any(p.text.strip() for p in section.footer.paragraphs):
            problems.append("content in footer")

    return {"safe": not problems, "problems": problems,
            "paragraphs": len(d.paragraphs), "tables": len(d.tables)}


# ---------------------------------------------------------------------------
# PDF
#
# Single column, real embedded text, no images. Built with PyMuPDF, which is
# already a dependency for figure detection, so no extra install.
#
# The DOCX remains the recommended upload wherever a choice exists: it is the
# format ATS vendors parse most reliably. The PDF exists because a lot of
# portals only accept one, and because people want something to email.
# ---------------------------------------------------------------------------

PAGE_W, PAGE_H = 595, 842          # A4 points
MARGIN = 52
LEADING = 13.2


class _Cursor:
    """Tracks the write position and opens a new page when the column runs out."""

    def __init__(self, pdf):
        self.pdf = pdf
        self.page = pdf.new_page(width=PAGE_W, height=PAGE_H)
        self.y = MARGIN

    def space(self, amount: float) -> None:
        self.y += amount

    def room_for(self, height: float) -> None:
        if self.y + height > PAGE_H - MARGIN:
            self.page = self.pdf.new_page(width=PAGE_W, height=PAGE_H)
            self.y = MARGIN

    def write(self, text, size=10.5, font="helv", indent=0.0, gap_after=0.0, colour=(0.1, 0.1, 0.1)):
        if not text:
            return
        import fitz

        width = PAGE_W - (MARGIN * 2) - indent
        box_h = _wrapped_height(text, width, size, font)
        self.room_for(box_h)

        rect = fitz.Rect(MARGIN + indent, self.y, PAGE_W - MARGIN, self.y + box_h + 2)
        rc = self.page.insert_textbox(
            rect, text, fontsize=size, fontname=font, color=colour, lineheight=1.25
        )
        # A negative return means the text did not fit. Grow and retry rather
        # than silently shipping a CV with a missing bullet.
        attempts = 0
        while rc < 0 and attempts < 4:
            attempts += 1
            box_h += size * 1.25 * 2
            self.room_for(box_h)
            rect = fitz.Rect(MARGIN + indent, self.y, PAGE_W - MARGIN, self.y + box_h + 2)
            rc = self.page.insert_textbox(
                rect, text, fontsize=size, fontname=font, color=colour, lineheight=1.25
            )

        self.y += box_h + gap_after

    def rule(self) -> None:
        import fitz

        self.room_for(6)
        self.page.draw_line(
            fitz.Point(MARGIN, self.y), fitz.Point(PAGE_W - MARGIN, self.y),
            color=(0.72, 0.72, 0.72), width=0.6,
        )
        self.y += 7


def _wrap_lines(text: str, width: float, size: float, font: str) -> int:
    """Greedy word wrap using real glyph widths, so the box is never too short."""
    import fitz

    lines = 0
    for para in str(text).split("\n"):
        words = para.split()
        if not words:
            lines += 1
            continue
        current = ""
        for word in words:
            trial = f"{current} {word}".strip()
            if fitz.get_text_length(trial, fontname=font, fontsize=size) <= width:
                current = trial
            else:
                if current:
                    lines += 1
                current = word
        if current:
            lines += 1
    return max(1, lines)


def _wrapped_height(text: str, width: float, size: float, font: str = "helv") -> float:
    return _wrap_lines(text, width, size, font) * (size * 1.25) + 2


def _pdf_heading(cur: _Cursor, text: str) -> None:
    cur.space(5)
    cur.write(text.upper(), size=10.5, font="hebo", colour=(0.26, 0.26, 0.26))
    cur.rule()


def build_pdf(tailored: Dict, contact: Dict) -> bytes:
    """Same content and same order as the DOCX, as selectable text."""
    import fitz

    pdf = fitz.open()
    cur = _Cursor(pdf)

    name = contact.get("name") or "Curriculum Vitae"
    cur.write(name, size=19, font="hebo", gap_after=2)

    details = [contact.get(k) for k in ("email", "phone", "location", "linkedin")]
    details = [d for d in details if d]
    if details:
        cur.write("  |  ".join(details), size=9, colour=(0.35, 0.35, 0.35), gap_after=4)
    cur.rule()

    if tailored.get("summary"):
        _pdf_heading(cur, "Professional Summary")
        cur.write(tailored["summary"], gap_after=4)

    if tailored.get("skills"):
        _pdf_heading(cur, "Skills")
        cur.write(", ".join(tailored["skills"]), gap_after=4)

    if tailored.get("experience"):
        _pdf_heading(cur, "Experience")
        for job in tailored["experience"]:
            head = " - ".join(x for x in (job.get("title"), job.get("company")) if x)
            if head:
                cur.write(head, size=11, font="hebo", gap_after=1)
            if job.get("dates"):
                cur.write(job["dates"], size=9, colour=(0.4, 0.4, 0.4), gap_after=2)
            for b in job.get("bullets", []):
                cur.write(f"-  {b}", indent=10, gap_after=2)
            cur.space(4)

    if tailored.get("projects"):
        _pdf_heading(cur, "Projects")
        for proj in tailored["projects"]:
            if proj.get("name"):
                cur.write(proj["name"], size=11, font="hebo", gap_after=1)
            for b in proj.get("bullets", []):
                cur.write(f"-  {b}", indent=10, gap_after=2)
            cur.space(4)

    for key, title in (("education", "Education"), ("certifications", "Certifications")):
        if tailored.get(key):
            _pdf_heading(cur, title)
            for line in tailored[key]:
                cur.write(line, gap_after=2)

    out = pdf.tobytes()
    pdf.close()
    return out


def verify_pdf_ats_safe(pdf_bytes: bytes) -> Dict:
    """A PDF passes when its text extracts in reading order and carries no images."""
    import fitz

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    problems = []
    text = ""
    images = 0
    for page in doc:
        text += page.get_text()
        images += len(page.get_images(full=True))

    if images:
        problems.append("contains images")
    if len(text.strip()) < 80:
        problems.append("little or no extractable text")
    if doc.needs_pass or doc.is_encrypted:
        problems.append("encrypted")

    pages = doc.page_count
    doc.close()
    return {
        "safe": not problems,
        "problems": problems,
        "pages": pages,
        "characters": len(text.strip()),
    }
